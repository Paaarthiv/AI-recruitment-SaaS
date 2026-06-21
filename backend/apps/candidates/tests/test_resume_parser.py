from io import BytesIO
from unittest.mock import patch

import docx
import httpx

from apps.candidates.resume_parser import (
    parse_json_object,
    parse_resume_text,
    request_ollama_parse,
    validate_parsed_payload,
)
from apps.candidates.tasks import extract_text_from_bytes


class FakeOllamaResponse:
    def __init__(self, response_text: str, eval_count: int = 1):
        self.response_text = response_text
        self.eval_count = eval_count

    def raise_for_status(self):
        return None

    def json(self):
        return {
            "response": self.response_text,
            "eval_count": self.eval_count,
            "done_reason": "stop",
        }


def make_pdf(text: str) -> bytes:
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("latin-1")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n"
        + stream + b"\nendstream",
    ]
    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf.extend(f"{index} 0 obj\n".encode("ascii"))
        pdf.extend(obj)
        pdf.extend(b"\nendobj\n")
    xref_offset = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    pdf.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    pdf.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    return bytes(pdf)


def make_docx(text: str) -> bytes:
    document = docx.Document()
    for line in text.splitlines():
        document.add_paragraph(line)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def make_table_docx(rows: list[tuple[str, str]]) -> bytes:
    document = docx.Document()
    table = document.add_table(rows=0, cols=2)
    for left, right in rows:
        row = table.add_row()
        row.cells[0].text = left
        row.cells[1].text = right
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_extracts_text_from_pdf_bytes():
    text = "Asha Patel asha@example.com Python Django"

    extracted = extract_text_from_bytes(make_pdf(text), "application/pdf")

    assert "Asha Patel" in extracted
    assert "Python Django" in extracted


def test_extracts_text_from_docx_bytes():
    text = "Asha Patel\nasha@example.com\nPython Django"

    extracted = extract_text_from_bytes(
        make_docx(text),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert "Asha Patel" in extracted
    assert "Python Django" in extracted


def test_extracts_text_from_docx_tables():
    extracted = extract_text_from_bytes(
        make_table_docx(
            [
                ("Parthiv A M", "amparthiv94@gmail.com"),
                ("Technical Skills", "Python Django React"),
            ]
        ),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert "Parthiv A M" in extracted
    assert "amparthiv94@gmail.com" in extracted
    assert "Python Django React" in extracted


def test_parse_json_object_extracts_wrapped_json():
    payload = parse_json_object(
        "Here is the JSON:\n"
        '{"personal_info": {"full_name": "Asha Patel"}, "skills": [],}\n'
        "Done"
    )

    assert payload == {"personal_info": {"full_name": "Asha Patel"}, "skills": []}


def test_parse_json_object_uses_last_valid_json_after_model_commentary():
    payload = parse_json_object(
        'The user asked for JSON like {ok: true}. '
        'The final answer is {"personal_info": {"full_name": "Asha Patel"}, "skills": []}'
    )

    assert payload == {"personal_info": {"full_name": "Asha Patel"}, "skills": []}


def test_ollama_request_accepts_json_after_explanatory_prefix(settings):
    settings.OLLAMA_BASE_URL = "http://ollama.test"
    settings.OLLAMA_API_KEY = ""
    settings.LLM_MODEL = "gpt-oss:20b"
    settings.OLLAMA_NUM_PREDICT = 4096
    settings.OLLAMA_TIMEOUT_SECONDS = 120
    response_text = (
        "I will return the structured resume JSON now.\n"
        '{"personal_info": {"full_name": "Asha Patel", "email": "asha@example.com"},'
        ' "summary": null, "skills": [], "experience": [], "education": [],'
        ' "certifications": [], "languages": [], "_metadata": {}}'
    )

    with patch(
        "apps.candidates.resume_parser.httpx.post",
        return_value=FakeOllamaResponse(response_text, eval_count=42),
    ):
        payload, usage = request_ollama_parse("Asha Patel\nasha@example.com", "gpt-oss:20b")

    assert payload["personal_info"]["email"] == "asha@example.com"
    assert usage == {"eval_count": 42}


def test_ollama_request_retries_malformed_json_with_larger_budget(settings):
    settings.OLLAMA_BASE_URL = "http://ollama.test"
    settings.OLLAMA_API_KEY = ""
    settings.LLM_MODEL = "gpt-oss:20b"
    settings.OLLAMA_NUM_PREDICT = 4096
    settings.OLLAMA_TIMEOUT_SECONDS = 120
    valid_response = (
        '{"personal_info": {"full_name": "Asha Patel", "email": "asha@example.com"},'
        ' "summary": null, "skills": [], "experience": [], "education": [],'
        ' "certifications": [], "languages": [], "_metadata": {}}'
    )
    calls = []

    def fake_post(*_args, **kwargs):
        calls.append(kwargs["json"]["options"]["num_predict"])
        if len(calls) == 1:
            return FakeOllamaResponse('{"personal_info": {"full_name": "Asha Patel"')
        return FakeOllamaResponse(valid_response, eval_count=42)

    with patch("apps.candidates.resume_parser.httpx.post", side_effect=fake_post):
        payload, usage = request_ollama_parse("Asha Patel\nasha@example.com", "gpt-oss:20b")

    assert calls == [4096, 8192]
    assert payload["personal_info"]["email"] == "asha@example.com"
    assert usage == {"eval_count": 42}


def test_ollama_parse_tolerates_real_resume_dates_and_estimates_confidence(settings):
    settings.LLM_MODEL = "gpt-oss:20b"
    payload = {
        "personal_info": {
            "full_name": "Asha Patel",
            "email": "asha@example.com",
            "phone": "+1 555 0101",
        },
        "summary": "Backend engineer focused on Django APIs.",
        "skills": [
            {"name": "Python", "proficiency": "advanced", "category": "programming_language"},
            {"name": "Django", "proficiency": "advanced", "category": "framework"},
            {"name": "PostgreSQL", "proficiency": "intermediate", "category": "database"},
            {"name": "Docker", "proficiency": "intermediate", "category": "devops"},
            {"name": "React", "proficiency": "beginner", "category": "framework"},
        ],
        "experience": [
            {
                "company": "Nexus Talent",
                "role": "Backend Engineer",
                "start_date": "2021",
                "end_date": "Present",
                "description": "Built recruiting APIs.",
                "achievements": ["Improved API latency"],
            }
        ],
        "education": [{"institution": "State University", "degree": "BS Computer Science"}],
        "_metadata": {"parsing_confidence": "low", "parsing_notes": []},
    }

    with patch(
        "apps.candidates.resume_parser.request_ollama_parse",
        return_value=(payload, {"eval_count": 100}),
    ):
        result = parse_resume_text(
            "Asha Patel\n"
            "asha@example.com | +1 555 0101\n"
            "Skills: Python, Django, PostgreSQL, Docker, React\n"
            "Nexus Talent - Backend Engineer - 2021 to Present\n"
            "State University - BS Computer Science"
        )

    assert result.model == "gpt-oss:20b"
    assert result.confidence == "high"
    assert result.data["experience"][0]["start_date"] == "2021-01"
    assert result.data["experience"][0]["end_date"] == "present"
    assert [skill["name"] for skill in result.data["skills"]] == [
        "Python",
        "Django",
        "PostgreSQL",
        "Docker",
        "React",
    ]


def test_parse_resume_tries_fallback_llm_before_heuristic(settings):
    settings.LLM_MODEL = "gpt-oss:20b"
    settings.LLM_FALLBACK_MODELS = ["gemma3:4b"]
    payload = {
        "personal_info": {
            "full_name": "Asha Patel",
            "email": "asha@example.com",
        },
        "summary": "Backend engineer focused on Django APIs.",
        "skills": [
            {"name": "Python", "proficiency": "advanced", "category": "programming_language"},
            {"name": "Django", "proficiency": "advanced", "category": "framework"},
        ],
        "experience": [],
        "education": [],
        "_metadata": {"parsing_confidence": "high", "parsing_notes": []},
    }
    calls = []

    def fake_request(_raw_text, model):
        calls.append(model)
        if model == "gpt-oss:20b":
            raise ValueError("Ollama returned empty response content.")
        return payload, {"eval_count": 80}

    with patch(
        "apps.candidates.resume_parser.request_ollama_parse",
        side_effect=fake_request,
    ):
        result = parse_resume_text("Asha Patel\nasha@example.com\nPython Django")

    assert calls == ["gpt-oss:20b", "gemma3:4b"]
    assert result.model == "gemma3:4b"
    assert result.validation_errors == []
    assert result.data["personal_info"]["email"] == "asha@example.com"


def test_ollama_parse_is_grounded_to_source_resume(settings):
    settings.LLM_MODEL = "gpt-oss:20b"
    payload = {
        "personal_info": {
            "full_name": "Wrong Person",
            "email": "wrong@example.com",
            "phone": "9999999999",
            "linkedin_url": "https://linkedin.com/in/wrong",
        },
        "summary": "Backend engineer.",
        "skills": [
            {"name": "Python", "proficiency": "advanced", "category": "programming_language"},
            {"name": "Kubernetes", "proficiency": "advanced", "category": "devops"},
        ],
        "experience": [
            {"company": "Imaginary Inc", "role": "CTO", "start_date": "2020"},
            {"company": "Nexus Talent", "role": "Backend Engineer", "start_date": "Jan 2021"},
        ],
        "education": [{"institution": "State University", "degree": "BS Computer Science"}],
        "_metadata": {"parsing_confidence": "high", "parsing_notes": []},
    }
    raw_text = (
        "Asha Patel\n"
        "asha@example.com | +1 555 0101 | linkedin.com/in/asha\n"
        "Skills: Python, Django, FastAPI, PostgreSQL\n"
        "Nexus Talent - Backend Engineer - Jan 2021 to Present\n"
        "State University - BS Computer Science"
    )

    with patch(
        "apps.candidates.resume_parser.request_ollama_parse",
        return_value=(payload, {"eval_count": 120}),
    ):
        result = parse_resume_text(raw_text)

    personal_info = result.data["personal_info"]
    assert personal_info["full_name"] == "Asha Patel"
    assert personal_info["email"] == "asha@example.com"
    assert personal_info["phone"] == "+1 555 0101"
    assert personal_info["linkedin_url"] == "https://linkedin.com/in/asha"
    assert "Kubernetes" not in [skill["name"] for skill in result.data["skills"]]
    assert [skill["name"] for skill in result.data["skills"]] == [
        "Python",
        "Django",
        "FastAPI",
        "PostgreSQL",
    ]
    assert result.data["experience"] == [
        {
            "company": "Nexus Talent",
            "role": "Backend Engineer",
            "start_date": "2021-01",
            "end_date": None,
            "location": None,
            "description": None,
            "achievements": [],
        }
    ]


def test_project_heavy_resume_is_grounded_without_fabrication(settings):
    settings.LLM_MODEL = "gpt-oss:20b"
    raw_text = (
        "Parthiv A M\n"
        "Full Stack Developer · Python | Django | React.js | Next.js · AI Integration\n"
        "Thiruvananthapuram · 9526500061 · amparthiv94@gmail.com · LinkedIn · GitHub\n"
        "Summary\n"
        "B.Tech CS graduate (2026) with hands-on experience building AI-integrated "
        "full-stack applications using Django, React, and Next.js.\n"
        "Technical Skills\n"
        "Languages: Python, JavaScript, TypeScript\n"
        "Frontend: React.js, Next.js, Tailwind CSS, HTML5, CSS3\n"
        "Backend: Django, Django REST Framework, FastAPI\n"
        "Databases: PostgreSQL, MySQL, Supabase\n"
        "AI / LLM: Ollama, Claude Code, OpenAI Codex, Gemini API\n"
        "DevOps & CI/CD: Docker, Kubernetes, GitHub Actions, Vercel, Railway\n"
        "Testing & Tools: Playwright, CodeRabbit, Testsprite, Postman, Git, GitHub, "
        "Firebase, IPFS\n"
        "Projects\n"
        "AI Recruitment SaaS Platform 2025 – 2026\n"
        "Django · Next.js · PostgreSQL · Supabase · Ollama · Docker · Kubernetes · "
        "GitHub Actions · Playwright\n"
        "Architected an enterprise-grade recruitment platform with secure auth.\n"
        "Decentralized Identity Management System 2024 – 2025\n"
        "Ethereum · Solidity · IPFS · MetaMask · Node.js · Express.js · Docker · "
        "REST APIs | github.com/Paaarthiv/MyDid_working\n"
        "Designed a blockchain-based identity system.\n"
        "Education\n"
        "B.Tech – Computer Science Engineering 2022 – 2026\n"
        "Lourdes Matha College of Science & Technology, Thiruvananthapuram · Coursework: DBMS\n"
        "Certifications\n"
        "Java Training — Spoken Tutorial Project, IIT Bombay September 2025\n"
    )
    weak_payload = {
        "personal_info": {
            "full_name": None,
            "email": "amparthiv94@gmail.com",
            "portfolio_url": "https://React.js",
        },
        "summary": None,
        "skills": [
            {"name": "React.js", "category": "framework"},
            {"name": "Node.js", "category": "framework"},
            {"name": "AWS", "category": "cloud"},
        ],
        "experience": [],
        "education": [
            {
                "institution": "University of California, Berkeley",
                "degree": "B.Sc. Computer Science",
            }
        ],
        "certifications": [],
        "languages": [{"language": "Python", "proficiency": "advanced"}],
        "_metadata": {"parsing_confidence": "high", "parsing_notes": ["Missing full name."]},
    }

    with patch(
        "apps.candidates.resume_parser.request_ollama_parse",
        return_value=(weak_payload, {"eval_count": 120}),
    ):
        result = parse_resume_text(raw_text)

    data = result.data
    skill_names = [skill["name"] for skill in data["skills"]]

    assert result.confidence == "high"
    assert data["personal_info"]["full_name"] == "Parthiv A M"
    assert data["personal_info"]["portfolio_url"] is None
    assert "AWS" not in skill_names
    assert "React" in skill_names
    assert "React.js" not in skill_names
    assert "Django REST Framework" in skill_names
    assert [project["name"] for project in data["projects"]] == [
        "AI Recruitment SaaS Platform",
        "Decentralized Identity Management System",
    ]
    assert data["education"] == [
        {
            "institution": "Lourdes Matha College of Science & Technology, Thiruvananthapuram",
            "degree": "B.Tech",
            "field_of_study": "Computer Science Engineering",
            "graduation_year": 2026,
            "gpa": None,
        }
    ]
    assert data["certifications"] == [
        {
            "name": "Java Training",
            "issuer": "Spoken Tutorial Project, IIT Bombay",
            "year": 2025,
            "credential_id": None,
        }
    ]
    assert data["languages"] == []
    assert "Missing full name." not in data["_metadata"]["parsing_notes"]


def test_validate_payload_tolerates_null_subfields():
    """LLMs emit ``null`` for sub-fields a resume omits (e.g. a language with no
    stated proficiency). These must NOT raise ValidationError — otherwise a single
    null discards an otherwise-good parse and forces the fallback model/heuristic.
    """
    payload = {
        "personal_info": {"full_name": "Asha Patel", "email": "asha@example.com"},
        "skills": [{"name": "Python", "proficiency": None, "category": None}],
        "experience": [{"company": "Nexus Talent", "role": None}],
        "education": [{"institution": None, "degree": "B.Tech"}],
        "certifications": [{"name": "AWS SAA", "issuer": None}],
        "languages": [
            {"language": "English", "proficiency": None},
            {"language": "Hindi", "proficiency": None},
        ],
        "_metadata": {},
    }

    validated = validate_parsed_payload(payload)

    assert validated.languages[0].language == "English"
    assert validated.languages[0].proficiency == ""
    assert validated.experience[0].company == "Nexus Talent"
    assert validated.experience[0].role == ""
    assert validated.skills[0].proficiency == "intermediate"
    assert validated.skills[0].category == "other"


def test_parse_resume_succeeds_on_primary_when_payload_has_null_subfields(settings):
    """End-to-end guard: a primary-model payload containing null sub-fields should
    be accepted on the FIRST model, never falling through to the fallback.
    """
    settings.LLM_MODEL = "qwen3-coder:480b"
    settings.LLM_FALLBACK_MODELS = ["gemma3:12b"]
    payload = {
        "personal_info": {"full_name": "Asha Patel", "email": "asha@example.com"},
        "summary": "Backend engineer.",
        "skills": [
            {"name": "Python", "proficiency": "advanced", "category": "programming_language"},
            {"name": "Django", "proficiency": None, "category": None},
        ],
        "experience": [],
        "education": [],
        "languages": [{"language": "English", "proficiency": None}],
        "_metadata": {"parsing_confidence": "high", "parsing_notes": []},
    }
    calls = []

    def fake_request(_raw_text, model):
        calls.append(model)
        return payload, {"eval_count": 90}

    with patch(
        "apps.candidates.resume_parser.request_ollama_parse",
        side_effect=fake_request,
    ):
        result = parse_resume_text(
            "Asha Patel\nasha@example.com\nSkills: Python, Django\nLanguages: English"
        )

    assert calls == ["qwen3-coder:480b"]
    assert result.model == "qwen3-coder:480b"
    assert result.data["personal_info"]["email"] == "asha@example.com"


def test_ollama_unavailable_uses_heuristic_fallback(settings):
    settings.LLM_MODEL = "gpt-oss:20b"

    with patch(
        "apps.candidates.resume_parser.request_ollama_parse",
        side_effect=httpx.ConnectError("connection refused"),
    ):
        result = parse_resume_text("Asha Patel\nasha@example.com\nPython Django")

    assert result.model == "heuristic-fallback"
    assert result.confidence == "medium"
    assert result.data["personal_info"]["email"] == "asha@example.com"
